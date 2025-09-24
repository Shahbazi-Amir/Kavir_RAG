# Lightweight loaders for TXT / PDF / DOCX / CSV.
# English comments inside code.

from pathlib import Path
from typing import List
import io

def load_txt_bytes(b: bytes, encoding: str = "utf-8") -> str:
    try:
        return b.decode(encoding)
    except UnicodeDecodeError:
        return b.decode("latin-1", errors="ignore")




# src/loaders.py  — replace only load_pdf_bytes()

def load_pdf_bytes(b: bytes, ocr: bool = False) -> str:
    # 1) Try PyMuPDF (fitz) first — best for mixed/RTL text
    print(">>> load_pdf_bytes called, ocr=", ocr)   # DEBUG
    try:
        import fitz  # PyMuPDF
        text_parts = []
        doc = fitz.open(stream=b, filetype="pdf")
        for page in doc:
            text_parts.append(page.get_text("text"))
        doc.close()
        joined = "\n".join([t for t in text_parts if t and t.strip()]).strip()
        if joined:
            return joined
    except Exception:
        pass

    # 2) Fallback to pypdf
    try:
        import io, pypdf
        text_parts = []
        reader = pypdf.PdfReader(io.BytesIO(b))
        for page in reader.pages:
            t = page.extract_text() or ""
            if t.strip():
                text_parts.append(t)
        joined = "\n".join(text_parts).strip()
        if joined or not ocr:
            return joined
    except Exception:
        if not ocr:
            return ""

    # 3) Final fallback: OCR (slow)
    try:
        from pdf2image import convert_from_bytes
        import pytesseract
        images = convert_from_bytes(b)
        ocr_texts = [pytesseract.image_to_string(im, lang="fas+eng") for im in images]
        return "\n".join(ocr_texts)
    except Exception:
        return ""


def load_docx_bytes(b: bytes) -> str:
    from docx import Document  # python-docx
    f = io.BytesIO(b)
    doc = Document(f)
    parts = []
    for p in doc.paragraphs:
        parts.append(p.text or "")
    # include table text (optional but useful)
    for table in doc.tables:
        for row in table.rows:
            parts.append("\t".join(cell.text or "" for cell in row.cells))
    return "\n".join(parts)

def load_csv_bytes(b: bytes) -> str:
    # Convert CSV to line-wise text (comma-joined) via pandas
    import pandas as pd
    f = io.BytesIO(b)
    try:
        df = pd.read_csv(f)
    except Exception:
        f.seek(0)
        df = pd.read_csv(f, encoding="latin-1")
    lines = []
    for _, row in df.iterrows():
        vals = [str(x) if x is not None else "" for x in row.tolist()]
        lines.append(", ".join(vals))
    return "\n".join(lines)

def load_any_bytes(b: bytes, ext: str, ocr: bool = False) -> str:
    ext = (ext or "").lower().lstrip(".")
    if ext == "txt":
        return load_txt_bytes(b)
    if ext == "pdf":
        return load_pdf_bytes(b, ocr=ocr)
    if ext in ("docx",):
        return load_docx_bytes(b)
    if ext in ("csv",):
        return load_csv_bytes(b)
    # default try utf-8 text
    return load_txt_bytes(b)
